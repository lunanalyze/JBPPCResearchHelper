from __future__ import annotations

import base64
import ctypes
import ctypes.wintypes
import datetime as dt
import json
import os
import re
import shutil
import socket
import sys
import threading
import time
import urllib.error
import urllib.request
import uuid
import webbrowser
from copy import deepcopy
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
from docx.table import Table

import collector
import paths
import updater


APP_NAME = "캄보디아&베트남 조사연구 도우미"
HOST = "127.0.0.1"
PORT = 8766
KEY_PATH = paths.CONFIG_DIR / "openai_key.bin"
TEMPLATE_PATH = paths.RESOURCES_DIR / "ppc_report_template.docx"
PROMPT_PATH = paths.RESOURCES_DIR / "report_prompt.md"
FAVICON_PATH = paths.APP_DIR / "PPC.ico"

paths.ensure_app_dirs()
paths.copy_default_resource("ppc_report_template.docx")
paths.copy_default_resource("report_prompt.md")

STATE = {
    "items": [],
    "run_dir": "",
    "report_path": "",
    "status": "준비되었습니다.",
}
JOBS: dict[str, dict] = {}
JOBS_LOCK = threading.Lock()


HTML = r"""<!doctype html>
<html lang="ko">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>PPC 조사연구 도우미</title>
<link rel="icon" href="/favicon.ico" type="image/x-icon">
<style>
  :root {
    color-scheme: light;
    --bg: #f6f7f9;
    --panel: #ffffff;
    --text: #17202a;
    --muted: #6b7280;
    --line: #d8dee8;
    --brand: #1f6feb;
    --brand-dark: #1557bd;
    --ok: #0f7b3d;
  }
  * { box-sizing: border-box; }
  body {
    margin: 0;
    font-family: "Malgun Gothic", "Segoe UI", sans-serif;
    background: var(--bg);
    color: var(--text);
  }
  header {
    padding: 18px 24px 14px;
    background: var(--panel);
    border-bottom: 1px solid var(--line);
  }
  .brand {
    display: flex;
    align-items: center;
    gap: 12px;
  }
  .brand-mark {
    width: 38px;
    height: 38px;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    border-radius: 8px;
    background: #eef4ff;
    color: #1557bd;
    border: 1px solid #c7ddff;
    font-size: 16px;
    font-weight: 900;
  }
  .brand-copy { min-width: 0; }
  h1 {
    margin: 0;
    font-size: 21px;
    font-weight: 700;
    letter-spacing: 0;
  }
  main { padding: 18px 24px 28px; }
  .controls {
    display: grid;
    grid-template-columns: minmax(260px, 1fr);
    gap: 12px;
    align-items: stretch;
    margin-bottom: 14px;
  }
  .group {
    min-height: 148px;
    display: flex;
    flex-direction: column;
    background: var(--panel);
    border: 1px solid var(--line);
    border-radius: 8px;
    padding: 12px;
  }
  .group-head {
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 10px;
    margin-bottom: 10px;
    min-height: 38px;
  }
  .group-title {
    font-size: 13px;
    font-weight: 700;
  }
  .date-row {
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 8px;
  }
  .button-row {
    display: flex;
    align-items: center;
    justify-content: flex-start;
    flex-wrap: wrap;
    gap: 8px;
    margin-top: auto;
    padding-top: 10px;
  }
  label {
    display: block;
    font-size: 12px;
    color: var(--muted);
  }
  input, select {
    width: 100%;
    height: 38px;
    margin-top: 4px;
    padding: 0 9px;
    border: 1px solid var(--line);
    border-radius: 6px;
    font: inherit;
    font-size: 13px;
    background: #fff;
    min-width: 0;
  }
  button {
    height: 38px;
    border: 1px solid var(--brand);
    border-radius: 6px;
    background: var(--brand);
    color: #fff;
    padding: 0 12px;
    font: inherit;
    font-weight: 700;
    cursor: pointer;
    white-space: nowrap;
  }
  button:hover { background: var(--brand-dark); }
  button.jb {
    background: #eef4ff;
    color: #1557bd;
    border-color: #c7ddff;
  }
  button.jb:hover {
    background: #dceaff;
    border-color: #9cc5ff;
  }
  button.ghost {
    background: #fff;
    color: var(--text);
    border-color: var(--line);
  }
  button.ghost:hover { background: #eef2f7; }
  button.secondary {
    background: var(--brand);
    color: #fff;
    border-color: var(--brand);
  }
  button.secondary:hover { background: var(--brand-dark); }
  button:disabled {
    opacity: 0.55;
    cursor: not-allowed;
  }
  input[type="checkbox"] {
    width: 16px;
    height: 16px;
    margin: 0;
    accent-color: var(--brand);
    cursor: pointer;
  }
  .status {
    min-height: 28px;
    margin: 8px 0 12px;
    color: var(--muted);
    font-size: 13px;
    white-space: pre-line;
  }
  .toolbar {
    display: flex;
    gap: 8px;
    align-items: center;
    justify-content: space-between;
    margin: 10px 0;
  }
  .summary {
    color: var(--muted);
    font-size: 13px;
  }
  .toolbar-actions {
    display: flex;
    gap: 8px;
    align-items: center;
    justify-content: flex-end;
    flex-wrap: wrap;
  }
  .side {
    display: flex;
    justify-content: flex-end;
    margin: -2px 0 14px;
  }
  .side button {
    min-width: 104px;
    padding: 0 12px;
  }
  .table-wrap {
    background: var(--panel);
    border: 1px solid var(--line);
    border-radius: 8px;
    overflow: auto;
    max-height: calc(100vh - 344px);
    cursor: default;
  }
  table {
    width: 100%;
    border-collapse: collapse;
    table-layout: fixed;
    font-size: 13px;
    cursor: default;
  }
  th, td {
    border-bottom: 1px solid var(--line);
    padding: 8px 10px;
    vertical-align: middle;
    white-space: nowrap;
    cursor: default;
  }
  th {
    position: sticky;
    top: 0;
    background: #f1f4f8;
    text-align: left;
    z-index: 1;
    font-weight: 800;
  }
  .sort-button {
    height: 26px;
    display: inline-flex;
    align-items: center;
    gap: 5px;
    padding: 0 7px;
    border: 1px solid transparent;
    background: transparent;
    color: var(--text);
    font: inherit;
    font-weight: 800;
  }
  .sort-button:hover,
  .sort-button.active {
    background: #fff;
    border-color: #cfd8e3;
  }
  .sort-indicator {
    width: 12px;
    color: var(--brand);
    font-size: 10px;
    line-height: 1;
  }
  .sort-menu {
    position: fixed;
    z-index: 30;
    width: 240px;
    padding: 6px;
    border: 1px solid var(--line);
    border-radius: 8px;
    background: #fff;
    box-shadow: 0 14px 32px rgba(15, 23, 42, 0.18);
  }
  .sort-menu button {
    width: 100%;
    height: 30px;
    display: flex;
    align-items: center;
    justify-content: flex-start;
    padding: 0 9px;
    border: 0;
    border-radius: 6px;
    background: #fff;
    color: var(--text);
    font-size: 12px;
    font-weight: 700;
  }
  .sort-menu button:hover {
    background: #eef4ff;
    color: #1557bd;
  }
  .sort-menu-divider {
    height: 1px;
    margin: 6px 2px;
    background: var(--line);
  }
  .filter-actions {
    display: grid;
    grid-template-columns: 1fr 1fr;
    gap: 6px;
    margin: 6px 0;
  }
  .filter-actions button {
    justify-content: center;
    border: 1px solid var(--line);
    font-size: 11px;
  }
  .filter-values {
    max-height: 220px;
    overflow: auto;
    padding: 4px 2px;
  }
  .filter-value {
    min-height: 28px;
    display: flex;
    align-items: center;
    gap: 7px;
    padding: 4px 6px;
    border-radius: 6px;
    color: var(--text);
    font-size: 12px;
    cursor: pointer;
  }
  .filter-value:hover { background: #f1f5fb; }
  .filter-value span {
    min-width: 0;
    overflow: hidden;
    text-overflow: ellipsis;
    white-space: nowrap;
  }
  tr:hover { background: #f7fbff; }
  .w-check { width: 50px; text-align: center; }
  .w-kind { width: 178px; }
  .w-date { width: 82px; }
  .w-source { width: 180px; }
  .w-format { width: 72px; }
  .w-open { width: 70px; text-align: center; }
  .section-select {
    height: 30px;
    margin: 0;
    padding: 0 28px 0 10px;
    border-color: #c7ddff;
    background: #eef4ff;
    color: #1557bd;
    font-size: 12px;
    font-weight: 800;
  }
  .format-badge {
    display: inline-flex;
    height: 24px;
    align-items: center;
    justify-content: center;
    padding: 0 8px;
    border-radius: 999px;
    background: #f1f4f8;
    color: #475467;
    font-size: 12px;
    font-weight: 800;
  }
  td.title {
    min-width: 420px;
    max-width: 760px;
    white-space: normal;
    line-height: 1.35;
    overflow-wrap: anywhere;
  }
  .muted {
    margin-top: 4px;
    color: var(--muted);
    font-size: 12px;
  }
  .open-item-btn {
    width: 40px;
    height: 30px;
    padding: 0;
    background: #fff;
    color: var(--text);
    border: 1px solid var(--line);
    font-size: 12px;
    font-weight: 800;
  }
  .open-item-btn:hover {
    background: #eef4ff;
    border-color: #9cc5ff;
    color: #1557bd;
  }
  .history {
    display: flex;
    gap: 8px;
    flex-wrap: wrap;
    margin-top: 12px;
  }
  .modal-backdrop {
    position: fixed;
    inset: 0;
    display: none;
    align-items: center;
    justify-content: center;
    padding: 24px;
    background: rgba(15, 23, 42, 0.42);
    z-index: 20;
  }
  .modal-backdrop.open { display: flex; }
  .modal {
    width: min(760px, 100%);
    max-height: min(720px, calc(100vh - 48px));
    display: grid;
    grid-template-rows: auto auto 1fr;
    background: #fff;
    border: 1px solid var(--line);
    border-radius: 8px;
    box-shadow: 0 24px 70px rgba(15, 23, 42, 0.28);
    overflow: hidden;
  }
  .modal-head {
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 12px;
    padding: 16px 18px;
    border-bottom: 1px solid var(--line);
  }
  .modal-title {
    font-size: 17px;
    font-weight: 800;
  }
  .modal-close {
    width: 34px;
    height: 34px;
    padding: 0;
    background: #fff;
    color: var(--text);
    border-color: var(--line);
    font-size: 18px;
    line-height: 1;
  }
  .modal-close:hover { background: #eef2f7; }
  .modal-help {
    padding: 10px 18px;
    color: var(--muted);
    font-size: 13px;
  }
  .modal-actions {
    display: flex;
    gap: 8px;
    justify-content: flex-end;
    padding: 0 18px 18px;
  }
  .update-bar {
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 10px 24px;
    background: #eef4ff;
    border-bottom: 1px solid #c7ddff;
    font-size: 13px;
  }
  .update-bar[hidden] { display: none; }
  .update-mark {
    width: 20px;
    height: 20px;
    flex: 0 0 auto;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    border-radius: 999px;
    background: var(--brand);
    color: #fff;
    font-size: 12px;
    font-weight: 900;
  }
  .update-copy {
    flex: 1 1 auto;
    min-width: 0;
    overflow: hidden;
    text-overflow: ellipsis;
    white-space: nowrap;
  }
  .update-copy strong { font-weight: 800; }
  .update-muted { color: var(--muted); }
  .update-dismiss {
    width: 28px;
    height: 28px;
    padding: 0;
    background: transparent;
    border-color: transparent;
    color: var(--muted);
    font-size: 16px;
    line-height: 1;
  }
  .update-dismiss:hover { background: #dbe7ff; }
  .update-modal {
    width: min(520px, 100%);
    grid-template-rows: auto auto minmax(0, 1fr);
  }
  .update-body {
    padding: 0 18px 18px;
    overflow: auto;
  }
  .update-notes {
    white-space: pre-wrap;
    max-height: 190px;
    overflow: auto;
    background: var(--bg);
    border: 1px solid var(--line);
    border-radius: 6px;
    padding: 10px 12px;
    font-size: 13px;
  }
  .update-notes[hidden] { display: none; }
  .update-guide {
    margin: 12px 0 0;
    color: var(--muted);
    font-size: 13px;
    line-height: 1.65;
  }
  .update-guide[hidden] { display: none; }
  .update-progress {
    display: flex;
    align-items: center;
    gap: 10px;
    margin-top: 12px;
    font-size: 13px;
  }
  .update-progress[hidden] { display: none; }
  .update-spinner {
    width: 16px;
    height: 16px;
    flex: 0 0 auto;
    border-radius: 999px;
    border: 2px solid #c7ddff;
    border-top-color: var(--brand);
    animation: update-spin 0.8s linear infinite;
  }
  @keyframes update-spin { to { transform: rotate(360deg); } }
  .update-error {
    margin-top: 12px;
    color: #b42318;
    font-size: 13px;
    white-space: pre-wrap;
  }
  .update-error[hidden] { display: none; }
  /* .modal-actions 가 display:flex 라 [hidden] 만으로는 안 숨는다. */
  .update-actions[hidden] { display: none; }
  .history-list {
    overflow: auto;
    padding: 8px;
  }
  .history-item {
    width: 100%;
    display: grid;
    grid-template-columns: 34px minmax(0, 1fr);
    gap: 12px;
    align-items: center;
    padding: 0;
    margin-bottom: 6px;
  }
  .history-bulk {
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 12px;
    padding: 10px 12px 0;
  }
  .history-bulk-left {
    display: grid;
    grid-template-columns: 34px auto;
    align-items: center;
    gap: 10px;
  }
  .history-check-label {
    display: inline-flex;
    align-items: center;
    color: var(--muted);
    font-size: 12px;
    font-weight: 800;
  }
  .history-check-box {
    width: 34px;
    height: 34px;
    display: flex;
    align-items: center;
    justify-content: center;
    background: #fff;
    border: 1px solid var(--line);
    border-radius: 7px;
    cursor: pointer;
  }
  .history-check-box:hover {
    border-color: var(--brand);
    background: #eef4ff;
  }
  .history-check-box input,
  .history-check {
    width: 16px;
    height: 16px;
    accent-color: var(--brand);
  }
  .history-card {
    min-height: 46px;
    display: grid;
    grid-template-columns: minmax(0, 1fr) auto;
    gap: 12px;
    align-items: center;
    padding: 12px 14px;
    background: #fff;
    color: var(--text);
    border: 1px solid var(--line);
    border-radius: 8px;
    text-align: left;
    font-weight: 600;
    cursor: pointer;
  }
  .history-card:hover {
    background: #f7fbff;
    border-color: #9cc5ff;
  }
  .history-actions {
    display: inline-flex;
    gap: 8px;
    align-items: center;
    justify-content: flex-end;
    flex-wrap: wrap;
  }
  .history-name {
    font-size: 14px;
    font-weight: 800;
  }
  .history-meta {
    margin-top: 4px;
    color: var(--muted);
    font-size: 12px;
    font-weight: 500;
  }
  .history-tags {
    display: inline-flex;
    gap: 6px;
    align-items: center;
    justify-content: flex-end;
    flex-wrap: wrap;
  }
  .history-tag {
    display: inline-flex;
    min-height: 24px;
    align-items: center;
    padding: 3px 8px;
    border-radius: 999px;
    background: #eef4ff;
    color: #1d4ed8;
    font-size: 12px;
    font-weight: 800;
  }
  .history-tag.report {
    background: #ecfdf3;
    color: #0f7b3d;
  }
  .history-delete {
    height: 30px;
    padding: 0 10px;
    background: #fff;
    color: #b42318;
    border: 1px solid #f1b9b4;
    border-radius: 6px;
    font-size: 12px;
    font-weight: 800;
  }
  .history-delete:hover {
    background: #fff1f0;
    border-color: #e0776e;
  }
  .empty-state {
    padding: 18px;
    color: var(--muted);
    font-size: 13px;
    text-align: center;
  }
  .history button {
    background: #fff;
    color: var(--text);
    border-color: var(--line);
  }
  .history button:hover {
    background: #f7fbff;
    border-color: #9cc5ff;
    color: #1557bd;
  }
  .tabs {
    display: flex;
    flex-wrap: wrap;
    gap: 6px;
    margin: 8px 0 10px;
  }
  .tab {
    height: 32px;
    padding: 0 11px;
    border: 1px solid var(--line);
    border-radius: 999px;
    background: #fff;
    color: var(--text);
    font-weight: 600;
    cursor: pointer;
  }
  .tab:hover {
    background: #eef4ff;
    border-color: #9cc5ff;
    color: #1557bd;
  }
  .tab.active {
    background: #e8f1ff;
    border-color: #9cc5ff;
    color: #1557bd;
  }
  .badge {
    display: inline-flex;
    min-width: 24px;
    height: 22px;
    align-items: center;
    justify-content: center;
    margin-left: 4px;
    border-radius: 999px;
    background: #eef4ff;
    color: #1d4ed8;
    font-size: 12px;
    font-weight: 700;
  }
  @media (max-width: 1024px) {
    .controls { grid-template-columns: 1fr; }
    .date-row { grid-template-columns: 1fr 1fr; }
    .table-wrap { max-height: none; }
  }
</style>
</head>
<body>
<header>
  <div class="brand">
    <div class="brand-mark">JB</div>
    <div class="brand-copy">
      <h1>PPC 조사연구 도우미</h1>
    </div>
  </div>
</header>
<div class="update-bar" id="updateBar" hidden>
  <span class="update-mark" aria-hidden="true">&uarr;</span>
  <div class="update-copy" id="updateBarCopy"></div>
  <button class="jb" id="updateNowBtn" type="button">지금 업데이트</button>
  <button class="update-dismiss" id="updateDismissBtn" type="button" title="이 버전 알림 닫기" aria-label="이 버전 알림 닫기">&times;</button>
</div>
<main>
  <section class="controls">
    <div class="group">
      <div class="group-head">
        <div class="group-title">수집 기간</div>
        <button class="jb" onclick="collectNow()">자료 수집</button>
      </div>
      <div class="date-row">
        <label>시작일<input id="startDate" type="date"></label>
        <label>종료일<input id="endDate" type="date"></label>
      </div>
      <div class="button-row">
        <label style="flex:1 1 150px">수집 대상<select id="target"><option value="all">전체</option><option value="cambodia">캄보디아</option><option value="vietnam">베트남</option></select></label>
        <label style="width:130px">사이트별 건수<select id="maxPerSource"><option value="10">10건</option><option value="20">20건</option><option value="30">30건</option><option value="50">50건</option><option value="100">100건</option><option value="0">모두</option></select></label>
      </div>
    </div>
  </section>
  <input id="reportTitle" type="hidden">
  <div class="status" id="status">준비되었습니다.</div>
  <div class="toolbar">
    <div class="summary" id="summary">선택 0건 / 전체 0건</div>
    <div class="toolbar-actions">
      <button class="ghost" onclick="resetAll()">초기화</button>
      <button class="ghost" onclick="loadRuns()">실행 기록</button>
      <button class="secondary" id="reportBtn" onclick="generateReport()" disabled>보고서 생성</button>
      <button class="ghost" onclick="openReport()">보고서 열기</button>
    </div>
  </div>
  <div class="history" id="history"></div>
  <div class="tabs" id="tabs"></div>
  <div class="table-wrap">
    <table>
      <thead><tr><th class="w-check"><input id="headCheck" type="checkbox" onchange="toggleAll(this.checked)"></th><th class="w-kind"><button class="sort-button" type="button" data-filter-key="category"><span>구분</span><span class="sort-indicator"></span></button></th><th class="w-source"><button class="sort-button" type="button" data-filter-key="source_name"><span>사이트</span><span class="sort-indicator"></span></button></th><th class="w-date"><button class="sort-button" type="button" data-filter-key="published_mm_dd"><span>일자</span><span class="sort-indicator"></span></button></th><th>제목</th><th class="w-format">형식</th><th class="w-open"></th></tr></thead>
      <tbody id="rows"></tbody>
    </table>
  </div>
  <div class="sort-menu" id="sortMenu" hidden>
    <div class="filter-actions">
      <button type="button" id="filterSelectAllBtn">전체 선택</button>
      <button type="button" id="filterClearAllBtn">전체 해제</button>
    </div>
    <div class="sort-menu-divider"></div>
    <div class="filter-values" id="filterValues"></div>
  </div>
</main>
<div class="modal-backdrop" id="historyModal" aria-hidden="true">
  <section class="modal" role="dialog" aria-modal="true" aria-labelledby="historyTitle">
    <div class="modal-head">
      <div class="modal-title" id="historyTitle">실행 기록</div>
      <button class="modal-close" id="historyCloseBtn" type="button" title="닫기" onclick="closeHistoryModal()">×</button>
    </div>
    <div class="modal-help">불러올 실행 기록을 선택하세요. 보고서가 생성된 기록은 별도로 표시됩니다.</div>
    <div class="history-bulk">
      <div class="history-bulk-left">
        <label class="history-check-box" title="전체 선택"><input id="historySelectAll" type="checkbox" onchange="toggleHistorySelectAll(this.checked)"></label>
        <span class="history-check-label">전체 선택</span>
      </div>
      <button class="history-delete" id="historyDeleteSelectedBtn" type="button" onclick="deleteSelectedHistoryRuns()" disabled>삭제</button>
    </div>
    <div class="history-list" id="historyList"></div>
  </section>
</div>
<div class="modal-backdrop" id="updateModal" aria-hidden="true">
  <section class="modal update-modal" role="dialog" aria-modal="true" aria-labelledby="updateTitle">
    <div class="modal-head">
      <div class="modal-title" id="updateTitle">업데이트</div>
      <button class="modal-close" id="updateCloseBtn" type="button" title="닫기">&times;</button>
    </div>
    <div class="modal-help" id="updateHelp"></div>
    <div class="update-body">
      <div class="update-notes" id="updateNotes" hidden></div>
      <p class="update-guide" id="updateGuide" hidden></p>
      <div class="update-progress" id="updateProgress" hidden>
        <span class="update-spinner" aria-hidden="true"></span>
        <span id="updatePhase"></span>
      </div>
      <div class="update-error" id="updateError" hidden></div>
      <div class="modal-actions update-actions" id="updateActions">
        <button class="ghost" id="updateLaterBtn" type="button">나중에</button>
        <button class="jb" id="updateApplyBtn" type="button">업데이트</button>
      </div>
    </div>
  </section>
</div>
<script>
const sections = ["캄보디아 금융/경제","캄보디아 정치/사회","베트남 금융/경제","베트남 정치/사회"];
let items = [];
let activeCollectJobId = "";
let activeGenerateJobId = "";
let activeTab = "전체";
let activeFilterKey = "";
let columnFilters = {};
let historyRuns = [];
let selectedHistoryRuns = new Set();

function isoDate(d){ return d.toISOString().slice(0,10); }
function initDates(){
  const end = new Date();
  const start = new Date(); start.setDate(end.getDate()-14);
  startDate.value = isoDate(start); endDate.value = isoDate(end);
}
function setStatus(text){ document.getElementById("status").textContent = text; }
function optionHtml(value, selected){ return `<option value="${value}" ${value===selected?"selected":""}>${value}</option>`; }
function render(){
  renderTabs();
  updateFilterHeaders();
  const tbody = document.getElementById("rows");
  const visible = visibleItems();
  tbody.innerHTML = visible.map(({it, idx}) => `
    <tr>
      <td><input type="checkbox" ${it.selected!==false?"checked":""} onchange="items[${idx}].selected=this.checked; updateReportButton()"></td>
      <td><select class="section-select" onchange="items[${idx}].category=this.value; render()">${sections.map(s=>optionHtml(s,it.category)).join("")}</select></td>
      <td>${escapeHtml(it.source_name || "")}</td>
      <td>${it.published_mm_dd || ""}</td>
      <td class="title">${escapeHtml(it.title || "")}<div class="muted">${escapeHtml(it.notes || "")}</div></td>
      <td><span class="format-badge">기사</span></td>
      <td><button class="open-item-btn" onclick="window.open(items[${idx}].url,'_blank')">열기</button></td>
    </tr>`).join("");
  updateReportButton();
  const headCheck = document.getElementById("headCheck");
  if(headCheck){
    const selectable = visible.map(({it}) => it);
    headCheck.checked = selectable.length > 0 && selectable.every(it => it.selected !== false);
    headCheck.indeterminate = selectable.some(it => it.selected !== false) && !headCheck.checked;
  }
  const summary = document.getElementById("summary");
  if(summary) summary.textContent = `선택 ${items.filter(x => x.selected !== false).length}건 / 전체 ${items.length}건`;
}
function visibleItems(){
  return items
    .map((it, idx) => ({it, idx}))
    .filter(({it}) => activeTab === "전체" || it.category === activeTab)
    .filter(({it}) => rowPassesFilters(it))
    .sort((a, b) => categorySortValue(a.it) - categorySortValue(b.it) || dateSortValue(a.it) - dateSortValue(b.it));
}
function categorySortValue(item){
  const idx = sections.indexOf(item.category || "");
  return idx >= 0 ? idx : 999;
}
function dateSortValue(item){
  const value = String(item.published_date || "");
  const time = value ? Date.parse(value) : NaN;
  return Number.isNaN(time) ? 8640000000000000 : time;
}
function rowPassesFilters(item){
  for(const [key, selected] of Object.entries(columnFilters)){
    if(!selected) continue;
    if(selected.length === 0) return false;
    if(!selected.includes(String(item[key] || ""))) return false;
  }
  return true;
}
function uniqueColumnValues(key){
  const values = [];
  const seen = new Set();
  for(const item of items){
    if(activeTab !== "전체" && item.category !== activeTab) continue;
    const value = String(item[key] || "");
    if(seen.has(value)) continue;
    seen.add(value);
    values.push(value);
  }
  return values.sort((a,b)=>a.localeCompare(b,"ko",{numeric:true,sensitivity:"base"}));
}
function selectedValuesForKey(key, values){
  return new Set(columnFilters[key] || values);
}
function setColumnFilter(key, selected, values){
  if(selected.size === values.length) delete columnFilters[key];
  else columnFilters[key] = [...selected];
}
function updateFilterHeaders(){
  document.querySelectorAll(".sort-button").forEach(button => {
    const key = button.dataset.filterKey || "";
    const active = Boolean(columnFilters[key]);
    button.classList.toggle("active", active);
    const indicator = button.querySelector(".sort-indicator");
    if(indicator) indicator.textContent = active ? "●" : "";
  });
}
function renderFilterValues(key){
  const list = document.getElementById("filterValues");
  const values = uniqueColumnValues(key);
  const selected = selectedValuesForKey(key, values);
  list.innerHTML = values.map(value => `
    <label class="filter-value" title="${escapeHtml(value || "(빈 값)") }">
      <input type="checkbox" value="${escapeHtml(value)}" ${selected.has(value) ? "checked" : ""}>
      <span>${escapeHtml(value || "(빈 값)")}</span>
    </label>`).join("");
  list.querySelectorAll("input").forEach(input => {
    input.addEventListener("change", () => {
      const current = selectedValuesForKey(key, values);
      if(input.checked) current.add(input.value);
      else current.delete(input.value);
      setColumnFilter(key, current, values);
      render();
      renderFilterValues(key);
    });
  });
}
function openFilterMenu(button){
  const menu = document.getElementById("sortMenu");
  activeFilterKey = button.dataset.filterKey || "";
  const rect = button.getBoundingClientRect();
  renderFilterValues(activeFilterKey);
  menu.style.left = `${Math.min(rect.left, window.innerWidth - 252)}px`;
  menu.style.top = `${rect.bottom + 4}px`;
  menu.hidden = false;
}
function closeFilterMenu(){
  const menu = document.getElementById("sortMenu");
  if(menu) menu.hidden = true;
  activeFilterKey = "";
}
function renderTabs(){
  const tabs = document.getElementById("tabs");
  if(!tabs) return;
  const labels = ["전체", ...sections];
  const counts = Object.fromEntries(labels.map(label => [label, 0]));
  counts["전체"] = items.length;
  for(const item of items){
    if(counts[item.category] !== undefined) counts[item.category] += 1;
  }
  tabs.innerHTML = labels.map(label => `<button class="tab ${activeTab===label?"active":""}" onclick="activeTab='${label}'; render()">${label} <span class="badge">${counts[label] || 0}</span></button>`).join("");
}
function toggleAll(checked){
  for(const item of items){
    if(activeTab === "전체" || item.category === activeTab) item.selected = checked;
  }
  render();
}
function updateReportButton(){
  document.getElementById("reportBtn").disabled = !items.some(x => x.selected !== false);
  const summary = document.getElementById("summary");
  if(summary) summary.textContent = `선택 ${items.filter(x => x.selected !== false).length}건 / 전체 ${items.length}건`;
}
function escapeHtml(s){ return String(s).replace(/[&<>"']/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c])); }
async function postJson(url, body){
  const res = await fetch(url, {method:"POST", headers:{"Content-Type":"application/json"}, body:JSON.stringify(body)});
  const data = await res.json();
  if(!res.ok || !data.ok) throw new Error(data.error || "request failed");
  return data;
}
async function collectNow(){
  try{
    activeCollectJobId = "";
    items = [];
    activeTab = "전체";
    render();
    setStatus("자료 수집 시작");
    const started = await postJson("/collect-start", {
      start_date:startDate.value, end_date:endDate.value,
      max_per_source:Number(maxPerSource.value), target:target.value
    });
    activeCollectJobId = started.job_id || "";
    await pollCollect(started.job_id);
  }catch(err){ setStatus(err.message); }
}
async function pollCollect(jobId){
  let done = false;
  while(!done){
    await new Promise(resolve => setTimeout(resolve, 700));
    if(activeCollectJobId !== jobId) return;
    const res = await fetch("/collect-status?id=" + encodeURIComponent(jobId));
    const data = await res.json();
    if(!data.ok) throw new Error(data.error || "수집 상태를 확인하지 못했습니다.");
    setStatus(data.message || "");
    if(data.done){
      done = true;
      if(activeCollectJobId !== jobId) return;
      activeCollectJobId = "";
      if(data.error) throw new Error(data.error);
      items = data.items || [];
      reportTitle.value = data.report_title || reportTitle.value;
      render();
      setStatus(`수집 완료: ${items.length}건`);
    }
  }
}
async function generateReport(){
  activeGenerateJobId = "";
  setStatus("보고서 생성 시작");
  try{
    const started = await postJson("/generate-report-start", {items, report_title:reportTitle.value});
    activeGenerateJobId = started.job_id || "";
    await pollGenerate(started.job_id);
  }catch(err){
    if(String(err.message || "").includes("OpenAI API key")){
      const key = prompt("OpenAI API key를 입력하세요. 현재 앱 전용 저장소에 암호화 저장됩니다.");
      if(key){
        try{
          await postJson("/api-key", {api_key:key});
          const started = await postJson("/generate-report-start", {items, report_title:reportTitle.value});
          activeGenerateJobId = started.job_id || "";
          await pollGenerate(started.job_id);
          return;
        }catch(inner){ setStatus(inner.message); return; }
      }
    }
    setStatus(err.message);
  }
}
async function pollGenerate(jobId){
  let done = false;
  while(!done){
    await new Promise(resolve => setTimeout(resolve, 700));
    if(activeGenerateJobId !== jobId) return;
    const res = await fetch("/generate-report-status?id=" + encodeURIComponent(jobId));
    const data = await res.json();
    if(!data.ok) throw new Error(data.error || "보고서 생성 상태를 확인하지 못했습니다.");
    setStatus(data.message || "");
    if(data.done){
      done = true;
      if(activeGenerateJobId !== jobId) return;
      activeGenerateJobId = "";
      if(data.error) throw new Error(data.error);
      setStatus("보고서 생성 완료: " + data.docx);
    }
  }
}
async function resetAll(){
  try{
    await postJson("/cancel-current", {collect_job_id: activeCollectJobId, generate_job_id: activeGenerateJobId});
  }catch(err){}
  activeCollectJobId = "";
  activeGenerateJobId = "";
  items = [];
  activeTab = "전체";
  render();
  setStatus("초기화했습니다.");
}
async function saveKey(){
  const key = prompt("OpenAI API key를 입력하세요. 현재 앱 전용 저장소에 암호화 저장됩니다.");
  if(!key) return;
  try{ await postJson("/api-key", {api_key:key}); setStatus("API key를 저장했습니다."); }
  catch(err){ setStatus(err.message); }
}
async function openReport(){ try{ const data = await postJson("/open-report", {}); setStatus(data.message || "보고서를 열었습니다."); }catch(err){ setStatus(err.message); } }
async function loadRuns(){
  const res = await fetch("/runs"); const data = await res.json();
  const modal = document.getElementById("historyModal");
  const el = document.getElementById("historyList");
  historyRuns = data.runs || [];
  selectedHistoryRuns = new Set();
  updateHistoryBulkControls();
  if(!historyRuns.length){
    el.innerHTML = `<div class="empty-state">저장된 실행 기록이 없습니다.</div>`;
  }else{
    el.innerHTML = historyRuns.map(r => `
      <div class="history-item">
        <label class="history-check-box" title="실행 기록 선택">
          <input class="history-check" type="checkbox" onchange="toggleHistoryRun('${r.id}', this.checked)" onclick="event.stopPropagation()">
        </label>
        <div class="history-card" onclick="loadRun('${r.id}'); closeHistoryModal()">
          <div>
            <div class="history-name">${escapeHtml(formatRunName(r.id))}</div>
            <div class="history-meta">${escapeHtml(r.id)}</div>
          </div>
          <div class="history-actions">
            <div class="history-tags">
              <span class="history-tag">${Number(r.item_count || 0)}건</span>
              ${r.has_report ? '<span class="history-tag report">보고서</span>' : ''}
            </div>
          </div>
        </div>
      </div>`).join("");
  }
  modal.classList.add("open");
  modal.setAttribute("aria-hidden","false");
}
function formatRunName(name){
  const m = String(name).match(/^(\d{2})(\d{2})(\d{2})_(\d{2})(\d{2})(?:_(\d+))?$/);
  if(!m) return name;
  const suffix = m[6] ? ` #${m[6]}` : "";
  return `20${m[1]}.${m[2]}.${m[3]} ${m[4]}:${m[5]}${suffix}`;
}
function toggleHistoryRun(id, checked){
  if(checked) selectedHistoryRuns.add(id);
  else selectedHistoryRuns.delete(id);
  updateHistoryBulkControls();
}
function toggleHistorySelectAll(checked){
  selectedHistoryRuns = checked ? new Set(historyRuns.map(r => r.id)) : new Set();
  document.querySelectorAll("#historyList .history-check").forEach(input => input.checked = checked);
  updateHistoryBulkControls();
}
function updateHistoryBulkControls(){
  const total = historyRuns.length;
  const selected = selectedHistoryRuns.size;
  const all = document.getElementById("historySelectAll");
  const deleteBtn = document.getElementById("historyDeleteSelectedBtn");
  if(all){
    all.checked = total > 0 && selected === total;
    all.indeterminate = selected > 0 && selected < total;
  }
  if(deleteBtn) deleteBtn.disabled = selected === 0;
}
async function deleteSelectedHistoryRuns(){
  const ids = [...selectedHistoryRuns];
  if(!ids.length) return;
  if(!confirm(`${ids.length}개 실행 기록을 삭제할까요?\n\n관련 metadata, JSON, 생성 보고서가 모두 삭제됩니다.`)) return;
  for(const id of ids){
    const res = await fetch("/delete-run", {
      method:"POST",
      headers:{"Content-Type":"application/json"},
      body:JSON.stringify({id})
    });
    const data = await res.json();
    if(!res.ok || !data.ok){
      setStatus("실행 기록 삭제 오류: " + (data.error || id));
      return;
    }
    if(data.cleared_current){
      items = [];
      render();
    }
  }
  setStatus(`${ids.length}개 실행 기록을 삭제했습니다.`);
  await loadRuns();
}
function closeHistoryModal(){
  const modal = document.getElementById("historyModal");
  modal.classList.remove("open");
  modal.setAttribute("aria-hidden","true");
}
async function loadRun(id){
  const res = await fetch("/run?id="+encodeURIComponent(id)); const data = await res.json();
  if(!data.ok){ setStatus(data.error || "실행 기록을 불러오지 못했습니다."); return; }
  items = data.items || []; reportTitle.value = data.report_title || reportTitle.value; render(); setStatus("실행 기록을 불러왔습니다.");
}
initDates();
document.querySelectorAll(".sort-button").forEach(button => button.addEventListener("click", () => openFilterMenu(button)));
document.getElementById("filterSelectAllBtn").addEventListener("click", () => {
  if(!activeFilterKey) return;
  delete columnFilters[activeFilterKey];
  render();
  renderFilterValues(activeFilterKey);
});
document.getElementById("filterClearAllBtn").addEventListener("click", () => {
  if(!activeFilterKey) return;
  columnFilters[activeFilterKey] = [];
  render();
  renderFilterValues(activeFilterKey);
});
document.addEventListener("click", event => {
  if(!event.target.closest("#sortMenu") && !event.target.closest(".sort-button")) closeFilterMenu();
  if(event.target.id === "historyModal") closeHistoryModal();
  if(event.target.id === "updateModal") closeUpdateModal();
});

/* ── 업데이트 ──────────────────────────────────────────────────────────
   배너 → 확인 창 → POST /update/apply. 서버는 팩을 받아 SHA-256 을 검증하고 업데이터를
   띄운 뒤 스스로 종료한다. 그때부터 이 화면이 /heartbeat 를 두드리며 기다리다, 서버가
   새 버전으로 응답하면 스스로 새로고침한다 — 사용자가 앱을 닫거나 설치 파일을 받을 일이 없다. */
let updateInfo = null;
let updateApplying = false;
const UPDATE_DISMISS_KEY = "ppcrh-update-dismissed";

function readUpdateDismissed(){
  try{ return localStorage.getItem(UPDATE_DISMISS_KEY) || ""; }catch(err){ return ""; }
}
function writeUpdateDismissed(version){
  try{ localStorage.setItem(UPDATE_DISMISS_KEY, version || ""); }catch(err){ /* 알림만 다시 뜰 뿐이다 */ }
}
function formatUpdateSize(bytes){
  /* 실제 팩은 수십 MB 지만, 반올림해서 "0MB" 로 보이는 것보다는 단위를 낮추는 편이 낫다. */
  const value = Number(bytes || 0);
  if(value <= 0) return "";
  if(value >= 10 * 1024 * 1024) return `${(value / 1024 / 1024).toFixed(0)}MB`;
  if(value >= 1024 * 1024) return `${(value / 1024 / 1024).toFixed(1)}MB`;
  return `${Math.max(1, Math.round(value / 1024))}KB`;
}
function renderUpdateBar(){
  const bar = document.getElementById("updateBar");
  const show = Boolean(updateInfo && updateInfo.available)
    && readUpdateDismissed() !== updateInfo.latest
    && !updateApplying;
  bar.hidden = !show;
  if(!show) return;
  const detail = [`현재 ${updateInfo.current || ""}`];
  const size = formatUpdateSize(updateInfo.size);
  if(size) detail.push(size);
  if(updateInfo.restart_required !== false) detail.push("적용 시 앱이 자동으로 재시작됩니다");
  const firstNote = String(updateInfo.notes || "").split("\n")[0].trim();
  if(firstNote) detail.push(firstNote);
  document.getElementById("updateBarCopy").innerHTML =
    `<strong>새 버전 ${escapeHtml(updateInfo.latest || "")}</strong>`
    + `<span class="update-muted"> 사용 가능 · ${escapeHtml(detail.join(" · "))}</span>`;
}
async function checkForUpdate(){
  /* 실패해도 조용하다 — /update/check 는 망이 막혀 있어도 200 + note 로 온다.
     첫 화면에 에러를 띄우지 않는 것이 이 API 의 규격이다. */
  try{
    const res = await fetch("/update/check", {cache:"no-store"});
    if(!res.ok) return;
    updateInfo = await res.json();
  }catch(err){ return; }
  renderUpdateBar();
}
function setUpdateModalMode(mode){
  /* mode: "confirm" | "progress" | "error" */
  document.getElementById("updateNotes").hidden = mode !== "confirm" || !String(updateInfo && updateInfo.notes || "").trim();
  document.getElementById("updateGuide").hidden = mode !== "confirm";
  document.getElementById("updateProgress").hidden = mode !== "progress";
  document.getElementById("updateError").hidden = mode !== "error";
  document.getElementById("updateActions").hidden = mode === "progress";
  document.getElementById("updateLaterBtn").textContent = mode === "error" ? "닫기" : "나중에";
  document.getElementById("updateApplyBtn").hidden = mode !== "confirm";
  /* 적용 중에는 닫지 못하게 한다 — 이때 다른 작업을 시작하면 재시작에 함께 끊긴다. */
  document.getElementById("updateCloseBtn").hidden = mode === "progress";
}
function openUpdateModal(){
  if(!updateInfo || !updateInfo.available) return;
  const size = formatUpdateSize(updateInfo.size);
  document.getElementById("updateTitle").textContent = `버전 ${updateInfo.latest} 로 업데이트`;
  document.getElementById("updateHelp").textContent =
    `현재 ${updateInfo.current} · 새 버전 ${updateInfo.latest}`
    + (updateInfo.released_at ? ` · ${updateInfo.released_at}` : "");
  document.getElementById("updateNotes").textContent = String(updateInfo.notes || "");
  document.getElementById("updateGuide").innerHTML =
    (size ? `내려받을 용량 ${escapeHtml(size)}. ` : "")
    + (updateInfo.restart_required !== false
        ? "적용하면 앱이 스스로 종료·갱신·재시작합니다. 진행 중인 작업이 있으면 업데이트가 거부됩니다.<br>"
        : "재시작 없이 반영됩니다.<br>")
    + "수집 자료·실행 기록·API 키는 그대로 유지됩니다. 실패하면 이전 버전으로 자동 복구합니다.";
  setUpdateModalMode("confirm");
  const modal = document.getElementById("updateModal");
  modal.classList.add("open");
  modal.setAttribute("aria-hidden", "false");
}
function closeUpdateModal(){
  if(updateApplying) return;
  const modal = document.getElementById("updateModal");
  modal.classList.remove("open");
  modal.setAttribute("aria-hidden", "true");
}
function showUpdatePhase(message){
  document.getElementById("updatePhase").textContent = message;
  setUpdateModalMode("progress");
}
function showUpdateError(message){
  updateApplying = false;
  document.getElementById("updateError").textContent = message;
  setUpdateModalMode("error");
  renderUpdateBar();
}
/* 서버가 다시 뜰 때까지 기다린다 — 종료·교체·재기동까지 수십 초 걸릴 수 있다.

   주의: /update/apply 는 응답을 먼저 내보내고 잠시 뒤에 프로세스를 끝낸다. 그래서 이 함수가
   도는 첫 1~2초 동안은 **아직 살아 있는 옛 서버**가 정상 응답한다. 그것을 '돌아왔다'로 치면
   교체가 끝나기도 전에 새로고침해 버린다. 서버가 한 번 끊긴 것을 보거나, 버전이 바뀐 것을
   확인한 뒤에만 돌아온 것으로 본다. */
async function waitForServer(previousVersion, timeoutMs = 180000){
  const deadline = Date.now() + timeoutMs;
  let sawDown = false;
  while(Date.now() < deadline){
    try{
      const res = await fetch("/heartbeat", {cache:"no-store"});
      if(res.ok){
        const data = await res.json();
        const version = String(data.version || "");
        if(sawDown || (version && version !== previousVersion)) return version;
      }else{
        sawDown = true;
      }
    }catch(err){
      sawDown = true;   /* 아직 안 떴다 — 끊긴 것을 봤다는 표시이기도 하다 */
    }
    await new Promise(resolve => setTimeout(resolve, 2000));
  }
  return null;
}
async function applyUpdate(){
  if(updateApplying) return;
  updateApplying = true;
  renderUpdateBar();
  showUpdatePhase("업데이트 파일을 받고 검증하는 중입니다… (수십 초)");
  let payload = null;
  try{
    const res = await fetch("/update/apply", {method:"POST"});
    payload = await res.json().catch(() => null);
    if(!res.ok || !payload || payload.ok === false){
      /* 409 = 진행 중인 작업이 있음, 502 = 다운로드·검증 실패. 사유를 그대로 보여준다. */
      showUpdateError((payload && payload.error) || `업데이트를 시작하지 못했습니다. (HTTP ${res.status})`);
      return;
    }
  }catch(err){
    showUpdateError("업데이트를 시작하지 못했습니다: " + err);
    return;
  }
  showUpdatePhase("앱을 다시 시작하는 중입니다… 이 창을 닫지 마세요.");
  const previous = (updateInfo && updateInfo.current) || "";
  const version = await waitForServer(previous);
  if(version === null){
    showUpdateError("앱이 다시 뜨지 않았습니다. 이전 버전으로 되돌렸을 수 있습니다 — 시작 메뉴에서 다시 실행해 주세요.");
    return;
  }
  if(version === previous){
    /* 되돌아왔다 — 업데이터가 새 버전을 못 띄워 이전 버전을 복원한 경우다. */
    showUpdateError(`새 버전이 적용되지 않아 이전 버전(${version})으로 되돌아갔습니다. 업데이트 창에 남은 사유를 확인해 주세요.`);
    return;
  }
  showUpdatePhase(`업데이트 완료 (${version}) — 화면을 새로 불러옵니다.`);
  writeUpdateDismissed("");
  setTimeout(() => window.location.reload(), 1200);
}

document.getElementById("updateNowBtn").addEventListener("click", openUpdateModal);
document.getElementById("updateApplyBtn").addEventListener("click", applyUpdate);
document.getElementById("updateCloseBtn").addEventListener("click", closeUpdateModal);
document.getElementById("updateLaterBtn").addEventListener("click", closeUpdateModal);
document.getElementById("updateDismissBtn").addEventListener("click", () => {
  writeUpdateDismissed(updateInfo && updateInfo.latest);
  renderUpdateBar();
});
checkForUpdate();
</script>
</body>
</html>"""


class DATA_BLOB(ctypes.Structure):
    _fields_ = [("cbData", ctypes.wintypes.DWORD), ("pbData", ctypes.POINTER(ctypes.c_char))]


def protect_with_dpapi(secret: str) -> bytes:
    data = secret.encode("utf-8")
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptProtectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData)
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def unprotect_with_dpapi(data: bytes) -> str:
    blob_in = DATA_BLOB(len(data), ctypes.cast(ctypes.create_string_buffer(data), ctypes.POINTER(ctypes.c_char)))
    blob_out = DATA_BLOB()
    if not ctypes.windll.crypt32.CryptUnprotectData(
        ctypes.byref(blob_in), None, None, None, None, 0, ctypes.byref(blob_out)
    ):
        raise ctypes.WinError()
    try:
        return ctypes.string_at(blob_out.pbData, blob_out.cbData).decode("utf-8")
    finally:
        ctypes.windll.kernel32.LocalFree(blob_out.pbData)


def save_api_key(api_key: str) -> None:
    paths.CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    try:
        KEY_PATH.write_bytes(b"dpapi:" + protect_with_dpapi(api_key))
    except Exception:
        KEY_PATH.write_bytes(b"b64:" + base64.b64encode(api_key.encode("utf-8")))


def load_api_key() -> str:
    if not KEY_PATH.exists():
        return ""
    data = KEY_PATH.read_bytes()
    if data.startswith(b"dpapi:"):
        return unprotect_with_dpapi(data[6:])
    if data.startswith(b"b64:"):
        return base64.b64decode(data[4:]).decode("utf-8")
    return ""


def openai_json_request(path: str, payload: dict, api_key: str, timeout: int = 240) -> dict:
    req = urllib.request.Request(
        f"https://api.openai.com/v1{path}",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
    )
    try:
        opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
        with opener.open(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")
        raise RuntimeError(f"OpenAI API error {exc.code}: {detail}") from exc
    except (TimeoutError, socket.timeout) as exc:
        raise TimeoutError("OpenAI API 응답 시간 초과") from exc
    except urllib.error.URLError as exc:
        reason = getattr(exc, "reason", None)
        if isinstance(reason, (TimeoutError, socket.timeout)) or "timed out" in str(reason).lower():
            raise TimeoutError("OpenAI API 응답 시간 초과") from exc
        raise


def extract_response_text(response: dict) -> str:
    if response.get("output_text"):
        return str(response["output_text"])
    parts = []
    for output in response.get("output", []):
        for content in output.get("content", []):
            if "text" in content:
                parts.append(str(content["text"]))
    return "\n".join(parts)


def extract_json_object(text: str) -> dict:
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            return json.loads(match.group(0))
        raise


def classify_with_openai(material: dict) -> str:
    api_key = load_api_key()
    if not api_key:
        return ""
    prompt = (
        "Classify this Cambodia/Vietnam news article into exactly one label: "
        "CAMBODIA_ECONOMY, CAMBODIA_POLITICS, VIETNAM_ECONOMY, VIETNAM_POLITICS, EXCLUDE. "
        "Economy includes banking, finance, macroeconomy, trade, investment, real estate, industry. "
        "Politics includes government, society, diplomacy, law, crime, education, health, culture. "
        "Return JSON only: {\"section\":\"...\"}."
    )
    response = openai_json_request("/responses", {
        "model": "gpt-4.1-mini",
        "input": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": json.dumps(material, ensure_ascii=False)},
        ],
        "text": {"format": {"type": "json_object"}},
    }, api_key, timeout=120)
    data = extract_json_object(extract_response_text(response))
    section = str(data.get("section", "")).strip()
    return section if section in collector.SECTION_LABELS else ""


def default_report_title(today: dt.date | None = None) -> str:
    today = today or dt.date.today()
    first = today.replace(day=1)
    week = ((today.day + first.weekday() - 1) // 7) + 1
    return f"캄보디아&베트남 조사·연구 [{today.year}년 {today.month}월 {week}주차]"


def row_for_client(item: collector.Item, selected: bool = True) -> dict:
    source_name = collector.display_source_name(item.source_name)
    return {
        "selected": selected,
        "category": item.category,
        "source_name": source_name,
        "title": collector.clean_article_title(item.title, source_name),
        "url": item.url,
        "published_date": collector.normalize_date(item.published_date) or item.published_date,
        "published_mm_dd": item.published_mm_dd,
        "notes": item.notes,
        "extra": item.extra,
    }


def item_from_client(raw: dict) -> collector.Item:
    source_name = str(raw.get("source_name") or "")
    return collector.Item(
        category=str(raw.get("category") or ""),
        source_name=source_name,
        title=collector.clean_article_title(str(raw.get("title") or ""), source_name),
        url=str(raw.get("url") or ""),
        published_date=str(raw.get("published_date") or ""),
        notes=str(raw.get("notes") or ""),
        extra=raw.get("extra") if isinstance(raw.get("extra"), dict) else {},
    )


def sort_items_by_date_asc(items: list[collector.Item]) -> list[collector.Item]:
    section_order = {
        "캄보디아 금융/경제": 0,
        "캄보디아 정치/사회": 1,
        "베트남 금융/경제": 2,
        "베트남 정치/사회": 3,
    }

    def key(pair):
        index, item = pair
        parsed = collector.parse_date(item.published_date)
        return (section_order.get(item.category, 999), parsed or dt.date.max, index)

    return [item for _, item in sorted(enumerate(items), key=key)]


def build_report_base(items: list[collector.Item]) -> dict:
    items = sort_items_by_date_asc(items)
    groups = {
        "CAMBODIA_ECONOMY_SECTION": [],
        "CAMBODIA_POLITICS_SECTION": [],
        "VIETNAM_ECONOMY_SECTION": [],
        "VIETNAM_POLITICS_SECTION": [],
    }
    category_to_key = {
        "캄보디아 금융/경제": ("CAMBODIA_ECONOMY_SECTION", "cambodia_economy", "1"),
        "캄보디아 정치/사회": ("CAMBODIA_POLITICS_SECTION", "cambodia_politics", "2"),
        "베트남 금융/경제": ("VIETNAM_ECONOMY_SECTION", "vietnam_economy", "3"),
        "베트남 정치/사회": ("VIETNAM_POLITICS_SECTION", "vietnam_politics", "4"),
    }
    counters = {key: 0 for key in groups}
    for item in items:
        section_key, id_prefix, section_no = category_to_key.get(item.category, category_to_key["캄보디아 금융/경제"])
        counters[section_key] += 1
        idx = counters[section_key]
        content = ""
        if isinstance(item.extra, dict):
            content = item.extra.get("article_text") or ""
        groups[section_key].append({
            "ITEM_ID": f"{id_prefix}_{idx}",
            "NO": f"{section_no}.{idx}",
            "SOURCE_NAME": clean_cell_value(item.source_name),
            "TITLE": clean_cell_value(item.title),
            "URL": clean_cell_value(item.url),
            "PUBLISHED_MM_DD": item.published_mm_dd,
            "_CONTENT_TEXT": clean_cell_value(content[:5000]),
            "SUMMARY_BULLET_1": "",
            "SUMMARY_BULLET_2": "",
            "SUMMARY_BULLET_3": "",
        })
    return {key: {"ITEMS": value} for key, value in groups.items()}


def iter_report_items(report: dict):
    for key in ["CAMBODIA_ECONOMY_SECTION", "CAMBODIA_POLITICS_SECTION", "VIETNAM_ECONOMY_SECTION", "VIETNAM_POLITICS_SECTION"]:
        yield from report.get(key, {}).get("ITEMS", [])


def build_bullet_request_context(items: list[dict]) -> str:
    materials = []
    for item in items:
        materials.append({
            "ITEM_ID": item["ITEM_ID"],
            "SECTION_NO": item["NO"].split(".")[0],
            "TITLE": item["TITLE"],
            "SOURCE_NAME": item["SOURCE_NAME"],
            "URL": item["URL"],
            "PUBLISHED_MM_DD": item["PUBLISHED_MM_DD"],
            "CONTENT_TEXT": item.get("_CONTENT_TEXT", ""),
        })
    return json.dumps(materials, ensure_ascii=False, indent=2)


def request_openai_bullets(api_key: str, prompt: str, report_items: list[dict]) -> tuple[list[dict], str]:
    expected_ids = [item["ITEM_ID"] for item in report_items]
    guard = (
        "\n\n[필수 검증]\n"
        f"- 입력 ITEM_ID는 총 {len(expected_ids)}개입니다: {', '.join(expected_ids)}\n"
        "- 출력 ITEMS 배열에는 위 ITEM_ID를 하나도 빠뜨리지 말고 정확히 같은 개수로 반환하세요.\n"
        "- 기사 내용이 부족해도 해당 ITEM_ID는 반드시 포함하고, 확인 가능한 범위에서 보수적으로 요약하세요.\n"
    )
    full_prompt = prompt.replace("{{INPUT_MATERIALS}}", build_bullet_request_context(report_items)) + guard
    response = openai_json_request("/responses", {
        "model": "gpt-4.1-mini",
        "input": [{"role": "user", "content": full_prompt}],
        "text": {"format": {"type": "json_object"}},
    }, api_key)
    output = extract_response_text(response)
    data = extract_json_object(output)
    return data.get("ITEMS", []), output


def merge_bullet_rows(report_items: list[dict], bullet_rows: list[dict]) -> set[str]:
    by_id = {item["ITEM_ID"]: item for item in report_items}
    updated: set[str] = set()
    for bullet in bullet_rows:
        item_id = str(bullet.get("ITEM_ID", ""))
        target = by_id.get(item_id)
        if not target:
            continue
        has_content = False
        for key in ["SUMMARY_BULLET_1", "SUMMARY_BULLET_2", "SUMMARY_BULLET_3"]:
            value = clean_cell_value(bullet.get(key, ""))
            target[key] = value
            has_content = has_content or bool(value)
        if has_content:
            updated.add(item_id)
    return updated


def missing_bullet_items(report_items: list[dict]) -> list[dict]:
    return [
        item for item in report_items
        if not any(clean_cell_value(item.get(key, "")) for key in ["SUMMARY_BULLET_1", "SUMMARY_BULLET_2", "SUMMARY_BULLET_3"])
    ]


def generate_report_json(items: list[collector.Item]) -> tuple[dict, str]:
    return generate_report_json_with_progress(items)


def generate_report_json_with_progress(
    items: list[collector.Item],
    progress=None,
    is_cancelled=None,
) -> tuple[dict, str]:
    api_key = load_api_key()
    if not api_key:
        raise RuntimeError("OpenAI API key가 저장되어 있지 않습니다.")
    prompt = PROMPT_PATH.read_text(encoding="utf-8")
    items = localize_english_titles(items, api_key, progress=progress, is_cancelled=is_cancelled)
    report = build_report_base(items)
    all_items = list(iter_report_items(report))
    outputs = []
    batch_size = 8
    total_batches = max((len(all_items) + batch_size - 1) // batch_size, 1)
    for batch_index, start in enumerate(range(0, len(all_items), batch_size), start=1):
        if is_cancelled and is_cancelled():
            raise RuntimeError("작업이 중지되었습니다.")
        batch = all_items[start:start + batch_size]
        if progress:
            progress(f"OpenAI 요약 생성 중 ({batch_index}/{total_batches}, {len(batch)}건)")
        bullets, raw_output = request_openai_bullets(api_key, prompt, batch)
        if is_cancelled and is_cancelled():
            raise RuntimeError("작업이 중지되었습니다.")
        outputs.append(raw_output)
        merge_bullet_rows(batch, bullets)
        missing = missing_bullet_items(batch)
        retry_count = 0
        while missing and retry_count < 2:
            retry_count += 1
            if is_cancelled and is_cancelled():
                raise RuntimeError("작업이 중지되었습니다.")
            if progress:
                progress(f"OpenAI 누락 요약 재요청 중 ({len(missing)}건)")
            bullets, raw_output = request_openai_bullets(api_key, prompt, missing)
            outputs.append(raw_output)
            merge_bullet_rows(missing, bullets)
            missing = missing_bullet_items(missing)
        if missing:
            missing_ids = ", ".join(item["ITEM_ID"] for item in missing)
            raise RuntimeError(f"OpenAI 요약 응답에서 누락된 항목이 있습니다: {missing_ids}")
    for item in all_items:
        item.pop("_CONTENT_TEXT", None)
        item.pop("ITEM_ID", None)
    return report, "\n\n".join(outputs)


def localize_english_titles(
    items: list[collector.Item],
    api_key: str,
    progress=None,
    is_cancelled=None,
) -> list[collector.Item]:
    targets = [(index, item.title) for index, item in enumerate(items) if needs_korean_title(item.title)]
    if not targets:
        return items
    if progress:
        progress("영문 제목 번역 중")
    translated: dict[int, str] = {}
    for start in range(0, len(targets), 20):
        if is_cancelled and is_cancelled():
            raise RuntimeError("작업이 중지되었습니다.")
        batch = targets[start:start + 20]
        payload = {
            "items": [{"index": index, "title": title} for index, title in batch],
            "rule": (
                "Translate each English news title into natural Korean. Preserve proper nouns and numbers. "
                "The Korean title must end as a headline noun or noun phrase. Avoid report-style or sentence-style "
                "endings such as ~임, ~함, ~했다, ~한다, ~됨. For example, use '완공' instead of '완공함', "
                "'확인' instead of '확인됨', and '예정' instead of '예정임'. Return JSON only."
            ),
        }
        response = openai_json_request("/responses", {
            "model": "gpt-4.1-mini",
            "input": [
                {
                    "role": "system",
                    "content": (
                        "Return JSON only: {\"items\":[{\"index\":0,\"korean_title\":\"...\"}]}. Do not add dates. "
                        "Every korean_title must be a Korean news headline that ends with a word/noun phrase, "
                        "not with endings like ~임, ~함, ~했다, ~한다, or ~됨."
                    ),
                },
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
            "text": {"format": {"type": "json_object"}},
        }, api_key, timeout=120)
        data = extract_json_object(extract_response_text(response))
        for row in data.get("items", []):
            try:
                index = int(row.get("index"))
            except Exception:
                continue
            korean = normalize_korean_headline_ending(clean_cell_value(row.get("korean_title", "")))
            if korean:
                translated[index] = korean
    out: list[collector.Item] = []
    for index, item in enumerate(items):
        korean = translated.get(index)
        if not korean:
            out.append(item)
            continue
        original = clean_cell_value(item.title)
        if original and original not in korean:
            item = collector.Item(
                category=item.category,
                source_name=item.source_name,
                title=f"{korean} ({original})",
                url=item.url,
                published_date=item.published_date,
                file_type=item.file_type,
                local_path=item.local_path,
                notes=item.notes,
                original_url=item.original_url,
                extra=item.extra,
            )
        out.append(item)
    return out


def normalize_korean_headline_ending(title: str) -> str:
    text = re.sub(r"\s+", " ", str(title or "")).strip(" \t\r\n.。!！?？")
    if not text:
        return ""
    replacements = {
        " 완료함": " 완료",
        " 완공함": " 완공",
        " 착공함": " 착공",
        " 개통함": " 개통",
        " 발표함": " 발표",
        " 밝힘": " 발표",
        " 확인됨": " 확인",
        " 승인함": " 승인",
        " 체결함": " 체결",
        " 개시함": " 개시",
        " 추진함": " 추진",
        " 확대함": " 확대",
        " 강화함": " 강화",
        " 증가함": " 증가",
        " 감소함": " 감소",
        " 상승함": " 상승",
        " 하락함": " 하락",
        " 개최함": " 개최",
        " 예정임": " 예정",
        " 전망임": " 전망",
        " 상황임": " 상황",
    }
    changed = True
    while changed:
        changed = False
        for ending, replacement in replacements.items():
            if text.endswith(ending):
                text = text[: -len(ending)] + replacement
                changed = True
    text = re.sub(r"(했다|하였다|한다|했습니다|하였습니다|합니다)$", "", text).strip()
    text = re.sub(r"(됨|됐음|되었음)$", "", text).strip()
    if text.endswith("함") and len(text) > 2:
        text = text[:-1].strip()
    return text.strip(" \t\r\n.。!！?？")


def needs_korean_title(title: str) -> bool:
    text = clean_cell_value(title)
    if not text:
        return False
    has_hangul = any("\uac00" <= ch <= "\ud7a3" for ch in text)
    has_alpha = any(("a" <= ch.lower() <= "z") for ch in text)
    return has_alpha and not has_hangul


def clean_cell_value(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def generate_docx(report_json: dict, output_path: Path, report_title: str) -> None:
    template = TEMPLATE_PATH
    doc = Document(template) if template.exists() else Document()
    set_report_title(doc, report_title)
    tables = doc.tables
    section_order = [
        ("CAMBODIA_ECONOMY_SECTION", "1. 캄보디아 금융/경제"),
        ("CAMBODIA_POLITICS_SECTION", "2. 캄보디아 정치/사회"),
        ("VIETNAM_ECONOMY_SECTION", "3. 베트남 금융/경제"),
        ("VIETNAM_POLITICS_SECTION", "4. 베트남 정치/사회"),
    ]
    for index, (key, label) in enumerate(section_order):
        items = report_json.get(key, {}).get("ITEMS", [])
        if len(tables) > 2 + index:
            fill_summary_table(tables[2 + index], items, label)
    detail_tables = list(doc.tables[6:10])
    for index in range(len(section_order) - 1, -1, -1):
        key, label = section_order[index]
        if index >= len(detail_tables):
            continue
        chunks = chunk_items(report_json.get(key, {}).get("ITEMS", []), 7)
        fill_detail_table_chunks(detail_tables[index], chunks, label)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = output_path.with_name(f".{output_path.stem}.{uuid.uuid4().hex[:8]}.tmp{output_path.suffix}")
    try:
        doc.save(tmp)
        replace_file_with_retry(tmp, output_path)
    finally:
        if tmp.exists():
            try:
                tmp.unlink()
            except Exception:
                pass


def set_report_title(doc: Document, report_title: str) -> None:
    if not report_title:
        return
    for table in doc.tables[:1]:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if "캄보디아&베트남 조사" in cell.text:
                    set_cell_text(cell, report_title)
                    if row_index == 1:
                        format_cell_text(cell, font_name="HY헤드라인M", size_pt=17, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    return
        if len(table.rows) > 1 and table.rows[1].cells:
            cell = table.rows[1].cells[0]
            set_cell_text(cell, report_title)
            format_cell_text(cell, font_name="HY헤드라인M", size_pt=17, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            return
    for paragraph in doc.paragraphs:
        if "캄보디아&베트남 조사" in paragraph.text:
            paragraph.text = report_title
            return


def replace_file_with_retry(src: Path, dst: Path, attempts: int = 5) -> None:
    last_exc: Exception | None = None
    for attempt in range(attempts):
        try:
            os.replace(src, dst)
            return
        except PermissionError as exc:
            last_exc = exc
            time.sleep(0.5 * (attempt + 1))
    raise PermissionError(f"{dst} 파일을 덮어쓸 수 없습니다. 보고서가 열려 있다면 닫은 뒤 다시 시도하세요.") from last_exc


def fill_summary_table(table, items: list[dict], label: str) -> None:
    desired = max(len(items), 1)
    remove_summary_header_row(table)
    data_start = summary_data_start_row(table)
    adjust_table_rows(table, data_start, desired)
    set_fixed_table_widths(table, [650, 9150, 1])
    if table.rows:
        for cell in unique_row_cells(table.rows[0]):
            set_cell_text(cell, summary_header_text(cell.text, label))
            format_cell_text(cell, font_name="맑은 고딕", size_pt=12, bold=True)
    rows = items + [blank_item()] * (desired - len(items))
    for row, item in zip(table.rows[data_start:], rows):
        cells = row.cells
        set_cell_text(cells[0], item.get("NO", ""))
        if len(cells) > 1:
            merge_cells(row, 1, len(cells) - 1)
            set_cell_text(row.cells[1], summary_title(item))


def remove_summary_header_row(table) -> None:
    for row in list(table.rows[1:3]):
        text = "\n".join(cell.text for cell in row.cells)
        if "순번" in text and "출처" in text and "기사제목" in text:
            delete_row(table, row)
            return


def summary_data_start_row(table) -> int:
    for idx, row in enumerate(table.rows):
        if any("{{NO}}" in cell.text for cell in row.cells):
            return idx
    return 1


def summary_header_text(template_text: str, label: str) -> str:
    match = re.search(r"\s*-+\s*$", str(template_text or ""))
    return f"{label}{match.group(0)}" if match else label


def unique_row_cells(row):
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        yield cell


def fill_detail_table(table, items: list[dict], label: str) -> None:
    desired = max(len(items), 1)
    adjust_detail_rows(table, desired)
    set_fixed_table_widths(table, [650, 9150])
    clear_table_row_heights(table)
    if table.rows:
        for cell in table.rows[0].cells:
            set_cell_text(cell, label)
            format_cell_text(cell, font_name="맑은 고딕", size_pt=12, bold=True)
    rows = items + [blank_item()] * (desired - len(items))
    for idx, item in enumerate(rows):
        title_row = table.rows[1 + idx * 3]
        url_row = table.rows[2 + idx * 3]
        bullet_row = table.rows[3 + idx * 3]
        number_cell = merge_vertical_cells(table, 1 + idx * 3, 3 + idx * 3, 0)
        set_cell_text(number_cell, item.get("NO", ""))
        format_cell_text(number_cell, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(title_row.cells[1], summary_title(item))
        set_cell_text(url_row.cells[1], item.get("URL", ""))
        set_cell_text(bullet_row.cells[1], detail_summary(item))
    clear_table_row_heights(table)


def fill_detail_table_chunks(base_table, chunks: list[list[dict]], label: str) -> None:
    if not chunks:
        chunks = [[]]
    template_xml = deepcopy(base_table._tbl)
    current = base_table
    for index, chunk in enumerate(chunks):
        if index:
            current = clone_table_after(current, template_xml)
        fill_detail_table(current, chunk, label)


def clone_table_after(table, table_xml) -> Table:
    new_tbl = deepcopy(table_xml)
    table._tbl.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def chunk_items(items: list[dict], size: int) -> list[list[dict]]:
    if not items:
        return [[]]
    return [items[index:index + size] for index in range(0, len(items), size)]


def adjust_table_rows(table, header_rows: int, desired_data_rows: int) -> None:
    while len(table.rows) < header_rows + desired_data_rows:
        clone_row(table, table.rows[-1])
    while len(table.rows) > header_rows + desired_data_rows:
        delete_row(table, table.rows[-1])


def adjust_detail_rows(table, desired_items: int) -> None:
    if len(table.rows) < 3:
        return
    title_row_xml = deepcopy(table.rows[1]._tr)
    url_row_xml = deepcopy(table.rows[2]._tr)
    bullet_row_xml = deepcopy(table.rows[2]._tr)
    remove_vertical_merges(title_row_xml)
    remove_vertical_merges(url_row_xml)
    remove_vertical_merges(bullet_row_xml)
    clear_row_height_xml(title_row_xml)
    clear_row_height_xml(url_row_xml)
    clear_row_height_xml(bullet_row_xml)
    while len(table.rows) > 1:
        delete_row(table, table.rows[-1])
    for _ in range(desired_items):
        table._tbl.append(deepcopy(title_row_xml))
        table._tbl.append(deepcopy(url_row_xml))
        table._tbl.append(deepcopy(bullet_row_xml))


def clone_row(table, row) -> None:
    table._tbl.append(deepcopy(row._tr))


def delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def merge_cells(row, start: int, end: int) -> None:
    if end <= start or start >= len(row.cells):
        return
    end = min(end, len(row.cells) - 1)
    try:
        if row.cells[start]._tc is not row.cells[end]._tc:
            row.cells[start].merge(row.cells[end])
    except Exception:
        pass


def merge_vertical_cells(table, start_row: int, end_row: int, col: int):
    start_row = max(start_row, 0)
    end_row = min(end_row, len(table.rows) - 1)
    if start_row >= end_row or col >= len(table.rows[start_row].cells):
        return table.rows[start_row].cells[col]
    try:
        return table.rows[start_row].cells[col].merge(table.rows[end_row].cells[col])
    except Exception:
        return table.rows[start_row].cells[col]


def remove_vertical_merges(row_xml) -> None:
    for tc_pr in row_xml.iter(qn("w:tcPr")):
        for v_merge in list(tc_pr.findall(qn("w:vMerge"))):
            tc_pr.remove(v_merge)


def clear_row_height_xml(row_xml) -> None:
    tr_pr = row_xml.find(qn("w:trPr"))
    if tr_pr is None:
        return
    for height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(height)


def clear_table_row_heights(table) -> None:
    for row in table.rows:
        clear_row_height_xml(row._tr)


def set_fixed_table_widths(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        seen = set()
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths) or id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


URL_RE = re.compile(r"https?://[^\s)]+")


def add_hyperlink(paragraph, url: str, text: str | None = None) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text or url
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_hyperlinks(paragraph, line: str) -> None:
    pos = 0
    for match in URL_RE.finditer(line):
        if match.start() > pos:
            paragraph.add_run(line[pos:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        pos = match.end()
    if pos < len(line):
        paragraph.add_run(line[pos:])


def set_cell_text(cell, text: str) -> None:
    text = str(text or "")
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    remove_paragraph_numbering(first)
    for child in list(first._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            first._p.remove(child)
    for idx, line in enumerate(text.split("\n")):
        if idx:
            first.add_run().add_break()
        add_text_with_hyperlinks(first, line)


def format_cell_text(
    cell,
    font_name: str | None = None,
    size_pt: int | float | None = None,
    bold: bool | None = None,
    alignment=None,
) -> None:
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        for run in paragraph.runs:
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if bold is not None:
                run.bold = bold


def remove_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def blank_item() -> dict:
    return {
        "NO": "",
        "SOURCE_NAME": "",
        "TITLE": "",
        "URL": "",
        "PUBLISHED_MM_DD": "",
        "SUMMARY_BULLET_1": "",
        "SUMMARY_BULLET_2": "",
        "SUMMARY_BULLET_3": "",
    }


def summary_title(item: dict) -> str:
    title = item.get("TITLE", "")
    date = item.get("PUBLISHED_MM_DD", "")
    return f"▣ {title} ({date})" if title else ""


def detail_title(item: dict) -> str:
    title = summary_title(item)
    url = item.get("URL", "")
    source = item.get("SOURCE_NAME", "")
    date = item.get("PUBLISHED_MM_DD", "")
    meta = " / ".join(x for x in [source, date] if x)
    if url and meta:
        return f"{title}\n{url} ({meta})"
    if url:
        return f"{title}\n{url}"
    return title


def detail_summary(item: dict) -> str:
    bullets = [
        item.get("SUMMARY_BULLET_1", ""),
        item.get("SUMMARY_BULLET_2", ""),
        item.get("SUMMARY_BULLET_3", ""),
    ]
    lines = []
    for bullet in bullets:
        text = clean_cell_value(bullet)
        if not text:
            continue
        lines.append(text if text.startswith("-") else f"- {text}")
    return "\n".join(lines)


def save_run_payload(run_dir: Path, report_title: str, items: list[collector.Item]) -> None:
    (run_dir / "run_config.json").write_text(
        json.dumps({"report_title": report_title}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    collector.write_outputs(items, run_dir)


def parse_max_per_source(value) -> int:
    try:
        number = int(value)
    except (TypeError, ValueError):
        return 10
    if number == 0:
        return 0  # 모두
    if number < 0:
        return 10
    return min(number, 500)


def update_collect_job(job_id: str, **updates) -> None:
    with JOBS_LOCK:
        job = JOBS.setdefault(job_id, {})
        for key, value in updates.items():
            if key == "message":
                job["message"] = str(value)
            else:
                job[key] = value


def is_job_cancelled(job_id: str) -> bool:
    with JOBS_LOCK:
        return bool((JOBS.get(job_id) or {}).get("cancelled"))


def run_collect_job(job_id: str, data: dict) -> None:
    try:
        end_date = parse_iso_date(data.get("end_date")) or dt.date.today()
        start_date = parse_iso_date(data.get("start_date")) or (end_date - dt.timedelta(days=14))
        target = str(data.get("target") or "all")
        # 0 = 모두(기간 안의 기사를 목록 끝까지). 그 외에는 사이트별 상한.
        max_per_source = parse_max_per_source(data.get("max_per_source"))
        update_collect_job(
            job_id,
            message="자료 수집 준비 중",
        )

        truncated: list[str] = []

        def progress(event_type: str, payload: dict) -> None:
            # 이 콜백은 기사 한 건마다 불린다. 취소를 여기서 확인해야 '모두'처럼
            # 오래 걸리는 수집도 중간에 멈출 수 있다.
            if is_job_cancelled(job_id):
                raise RuntimeError("작업이 중지되었습니다.")
            source = payload.get("source", "")
            if event_type == "source_truncated":
                if source not in truncated:
                    truncated.append(source)
            elif event_type == "source_start":
                update_collect_job(job_id, message=f"{source}: 수집 중")
            elif event_type == "source_page":
                update_collect_job(
                    job_id,
                    message=f"{source}: {payload.get('page', 1)}페이지 ({payload.get('count', 0)}건)",
                )
            elif event_type == "item_start":
                update_collect_job(job_id, message=f"{source}: {payload.get('count', 0)}건 수집됨")
            elif event_type == "source_done":
                update_collect_job(job_id, message=f"{source}: {payload.get('count', 0)}건 완료")

        items = collector.collect_all(
            start_date,
            end_date,
            max_per_source=max_per_source,
            include_cambodia=target in {"all", "cambodia"},
            include_vietnam=target in {"all", "vietnam"},
            classifier=classify_with_openai,
            progress=progress,
        )
        if is_job_cancelled(job_id):
            raise RuntimeError("작업이 중지되었습니다.")
        report_title = default_report_title(end_date)
        run_dir = collector.make_run_dir()
        if is_job_cancelled(job_id):
            raise RuntimeError("작업이 중지되었습니다.")
        save_run_payload(run_dir, report_title, items)
        if is_job_cancelled(job_id):
            try:
                shutil.rmtree(run_dir)
            except Exception:
                pass
            raise RuntimeError("작업이 중지되었습니다.")
        visible_items = [row_for_client(item) for item in items if item.category != "error"]
        # 목록 상한에서 끊긴 사이트가 있으면 완료 메시지에 그대로 적는다.
        truncated_note = (
            f" (목록 {collector.LIST_PAGE_LIMIT}페이지 상한에서 중단: {', '.join(truncated)})"
            if truncated else ""
        )
        STATE["items"] = items
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = ""
        update_collect_job(
            job_id,
            done=True,
            items=visible_items,
            report_title=report_title,
            run_dir=str(run_dir),
            message=f"수집 완료: {len(visible_items)}건{truncated_note}",
        )
    except Exception as exc:
        update_collect_job(job_id, done=True, error=str(exc), message=str(exc))


def run_generate_job(job_id: str, data: dict) -> None:
    try:
        raw_items = data.get("items") if isinstance(data.get("items"), list) else []
        selected = [item_from_client(raw) for raw in raw_items if raw.get("selected") is not False]
        if not selected:
            raise RuntimeError("보고서에 포함할 항목이 없습니다.")
        run_dir = Path(str(STATE.get("run_dir") or "")) if STATE.get("run_dir") else collector.make_run_dir()
        report_title = str(data.get("report_title") or default_report_title())
        update_collect_job(job_id, message="보고서 자료 정리 중", run_dir=str(run_dir))
        save_run_payload(run_dir, report_title, selected)

        report_json, raw_output = generate_report_json_with_progress(
            selected,
            progress=lambda message: update_collect_job(job_id, message=message),
            is_cancelled=lambda: is_job_cancelled(job_id),
        )
        if is_job_cancelled(job_id):
            raise RuntimeError("작업이 중지되었습니다.")
        update_collect_job(job_id, message="DOCX 생성 중")
        (run_dir / "report_data.json").write_text(json.dumps(report_json, ensure_ascii=False, indent=2), encoding="utf-8")
        (run_dir / "llm_output.txt").write_text(raw_output, encoding="utf-8")
        docx_path = run_dir / "generated_report.docx"
        generate_docx(report_json, docx_path, report_title)
        STATE["items"] = selected
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = str(docx_path)
        update_collect_job(job_id, done=True, docx=str(docx_path), message="보고서 생성 완료")
    except Exception as exc:
        update_collect_job(job_id, done=True, error=str(exc), message=str(exc))


class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt: str, *args) -> None:
        return

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/":
            self.send_text(HTML, "text/html; charset=utf-8")
        elif parsed.path == "/favicon.ico":
            self.handle_favicon()
        elif parsed.path == "/collect-status":
            self.handle_collect_status(parsed)
        elif parsed.path == "/generate-report-status":
            self.handle_generate_report_status(parsed)
        elif parsed.path == "/runs":
            self.handle_runs()
        elif parsed.path == "/run":
            self.handle_run(parsed)
        elif parsed.path == "/heartbeat":
            # version 을 함께 내려보낸다 — 업데이트 화면이 서버가 돌아오기를 기다리다가
            # 이 값으로 "새 버전이 떴다"를 확인한다.
            self.send_json({"ok": True, "version": updater.APP_VERSION})
        elif parsed.path == "/update/check":
            self.send_json(updater.check())
        elif parsed.path == "/update/version":
            self.send_json(
                {"ok": True, "version": updater.APP_VERSION, "installed": updater.installed()}
            )
        else:
            self.send_json({"ok": False, "error": "not found"}, 404)

    def do_POST(self) -> None:
        try:
            if self.path == "/collect-start":
                self.handle_collect_start()
            elif self.path == "/collect":
                self.handle_collect()
            elif self.path == "/generate-report-start":
                self.handle_generate_report_start()
            elif self.path == "/generate-report":
                self.handle_generate_report()
            elif self.path == "/cancel-current":
                self.handle_cancel_current()
            elif self.path == "/delete-run":
                self.handle_delete_run()
            elif self.path == "/api-key":
                data = self.read_json()
                save_api_key(str(data.get("api_key") or "").strip())
                self.send_json({"ok": True})
            elif self.path == "/open-report":
                self.handle_open_report()
            elif self.path == "/update/apply":
                self.handle_update_apply()
            elif self.path == "/heartbeat":
                self.send_json({"ok": True, "version": updater.APP_VERSION})
            else:
                self.send_json({"ok": False, "error": "not found"}, 404)
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, 500)

    def handle_favicon(self) -> None:
        if not FAVICON_PATH.exists():
            self.send_json({"ok": False, "error": "favicon not found"}, 404)
            return
        data = FAVICON_PATH.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", "image/x-icon")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def handle_collect_start(self) -> None:
        data = self.read_json()
        job_id = uuid.uuid4().hex
        with JOBS_LOCK:
            JOBS[job_id] = {
                "done": False,
                "error": "",
                "message": "자료 수집 준비 중",
                "cancelled": False,
                "items": [],
                "run_dir": "",
                "report_title": "",
            }
        thread = threading.Thread(target=run_collect_job, args=(job_id, data), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_collect_status(self, parsed) -> None:
        job_id = parse_qs(parsed.query).get("id", [""])[0]
        with JOBS_LOCK:
            job = dict(JOBS.get(job_id) or {})
        if not job:
            self.send_json({"ok": False, "error": "수집 작업을 찾을 수 없습니다."}, 404)
            return
        self.send_json({
            "ok": True,
            "done": bool(job.get("done")),
            "error": job.get("error", ""),
            "message": job.get("message", ""),
            "items": job.get("items", []),
            "run_dir": job.get("run_dir", ""),
            "report_title": job.get("report_title", ""),
        })

    def handle_generate_report_start(self) -> None:
        data = self.read_json()
        raw_items = data.get("items") if isinstance(data.get("items"), list) else []
        selected = [raw for raw in raw_items if raw.get("selected") is not False]
        if not selected:
            raise RuntimeError("보고서에 포함할 항목이 없습니다.")
        if not load_api_key():
            raise RuntimeError("OpenAI API key가 저장되어 있지 않습니다.")
        job_id = uuid.uuid4().hex
        with JOBS_LOCK:
            JOBS[job_id] = {
                "done": False,
                "error": "",
                "message": "보고서 생성 준비 중",
                "cancelled": False,
                "docx": "",
            }
        thread = threading.Thread(target=run_generate_job, args=(job_id, data), daemon=True)
        thread.start()
        self.send_json({"ok": True, "job_id": job_id})

    def handle_generate_report_status(self, parsed) -> None:
        job_id = parse_qs(parsed.query).get("id", [""])[0]
        with JOBS_LOCK:
            job = dict(JOBS.get(job_id) or {})
        if not job:
            self.send_json({"ok": False, "error": "보고서 생성 작업을 찾을 수 없습니다."}, 404)
            return
        self.send_json({
            "ok": True,
            "done": bool(job.get("done")),
            "error": job.get("error", ""),
            "message": job.get("message", ""),
            "docx": job.get("docx", ""),
        })

    def handle_cancel_current(self) -> None:
        data = self.read_json()
        job_ids = [str(data.get("collect_job_id") or ""), str(data.get("generate_job_id") or "")]
        with JOBS_LOCK:
            for job_id in job_ids:
                if job_id and job_id in JOBS and not JOBS[job_id].get("done"):
                    JOBS[job_id]["cancelled"] = True
                    JOBS[job_id]["message"] = "중지 요청됨"
        self.send_json({"ok": True})

    def handle_collect(self) -> None:
        data = self.read_json()
        end_date = parse_iso_date(data.get("end_date")) or dt.date.today()
        start_date = parse_iso_date(data.get("start_date")) or (end_date - dt.timedelta(days=14))
        target = str(data.get("target") or "all")
        # 0 = 모두(기간 안의 기사를 목록 끝까지). 그 외에는 사이트별 상한.
        max_per_source = parse_max_per_source(data.get("max_per_source"))
        run_dir = collector.make_run_dir()
        items = collector.collect_all(
            start_date,
            end_date,
            max_per_source=max_per_source,
            include_cambodia=target in {"all", "cambodia"},
            include_vietnam=target in {"all", "vietnam"},
            classifier=classify_with_openai,
        )
        report_title = default_report_title(end_date)
        save_run_payload(run_dir, report_title, items)
        STATE["items"] = items
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = ""
        self.send_json({
            "ok": True,
            "items": [row_for_client(item) for item in items if item.category != "error"],
            "run_dir": str(run_dir),
            "report_title": report_title,
        })

    def handle_generate_report(self) -> None:
        data = self.read_json()
        raw_items = data.get("items") if isinstance(data.get("items"), list) else []
        selected = [item_from_client(raw) for raw in raw_items if raw.get("selected") is not False]
        if not selected:
            raise RuntimeError("보고서에 포함할 항목이 없습니다.")
        run_dir = Path(str(STATE.get("run_dir") or "")) if STATE.get("run_dir") else collector.make_run_dir()
        report_title = str(data.get("report_title") or default_report_title())
        save_run_payload(run_dir, report_title, selected)
        report_json, raw_output = generate_report_json(selected)
        (run_dir / "report_data.json").write_text(json.dumps(report_json, ensure_ascii=False, indent=2), encoding="utf-8")
        (run_dir / "llm_output.txt").write_text(raw_output, encoding="utf-8")
        docx_path = run_dir / "generated_report.docx"
        generate_docx(report_json, docx_path, report_title)
        STATE["items"] = selected
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = str(docx_path)
        self.send_json({"ok": True, "docx": str(docx_path)})

    def handle_open_report(self) -> None:
        report_path = Path(str(STATE.get("report_path") or ""))
        if not report_path.exists() and STATE.get("run_dir"):
            candidate = Path(str(STATE["run_dir"])) / "generated_report.docx"
            if candidate.exists():
                report_path = candidate
        if not report_path.exists():
            raise RuntimeError("열 수 있는 보고서가 없습니다.")
        os.startfile(str(report_path))
        self.send_json({"ok": True, "message": str(report_path)})

    def handle_runs(self) -> None:
        runs = []
        if paths.RUNS_DIR.exists():
            for path in sorted((p for p in paths.RUNS_DIR.iterdir() if p.is_dir()), reverse=True):
                item_count = 0
                try:
                    item_count = len(collector.read_items(path))
                except Exception:
                    item_count = 0
                runs.append({"id": path.name, "has_report": (path / "generated_report.docx").exists(), "item_count": item_count})
        self.send_json({"ok": True, "runs": runs[:50]})

    def handle_run(self, parsed) -> None:
        run_id = parse_qs(parsed.query).get("id", [""])[0]
        run_dir = resolve_run_dir(run_id)
        if not run_dir:
            self.send_json({"ok": False, "error": "실행 기록을 찾을 수 없습니다."}, 404)
            return
        items = collector.read_items(run_dir)
        config = {}
        if (run_dir / "run_config.json").exists():
            config = json.loads((run_dir / "run_config.json").read_text(encoding="utf-8"))
        STATE["items"] = items
        STATE["run_dir"] = str(run_dir)
        STATE["report_path"] = str(run_dir / "generated_report.docx") if (run_dir / "generated_report.docx").exists() else ""
        self.send_json({
            "ok": True,
            "items": [row_for_client(item) for item in items],
            "report_title": config.get("report_title") or default_report_title(),
        })

    def handle_delete_run(self) -> None:
        data = self.read_json()
        run_id = str(data.get("id") or "")
        run_dir = resolve_run_dir(run_id)
        if not run_dir:
            raise RuntimeError("실행 기록을 찾을 수 없습니다.")
        runs_root = paths.RUNS_DIR.resolve()
        resolved = run_dir.resolve()
        try:
            resolved.relative_to(runs_root)
        except ValueError as exc:
            raise RuntimeError("실행 기록 폴더 밖은 삭제할 수 없습니다.") from exc
        cleared_current = False
        current_run = str(STATE.get("run_dir") or "")
        if current_run and same_path(current_run, resolved):
            STATE["items"] = []
            STATE["run_dir"] = ""
            STATE["report_path"] = ""
            cleared_current = True
        shutil.rmtree(resolved)
        self.send_json({"ok": True, "deleted": str(resolved), "cleared_current": cleared_current})

    def handle_update_apply(self) -> None:
        """업데이트 적용 — 받아서 검증하고 업데이터를 띄운 뒤 **이 프로세스를 종료**한다.

        스스로 죽는 이유: 실행 중인 exe 는 Windows 가 잠그고 있어 자기 자신을 갈아끼울 수 없다.
        """
        try:
            result = updater.apply(PORT, busy_reason=update_busy_reason)
        except updater.UpdateError as exc:
            self.send_json({"ok": False, "error": exc.message}, exc.status)
            return
        self.send_json(result)
        # 응답을 먼저 내보내고 종료한다 — 먼저 죽으면 화면은 "실패"로 본다.
        # os._exit 를 쓰는 것은 의도다. server.shutdown() 은 처리 중인 요청이 끝나기를 기다리고
        # 남은 스레드가 있으면 종료가 늘어지는데, 그동안 exe 잠금이 안 풀려 업데이터가 대기하다
        # 타임아웃된다. 어차피 저장할 상태는 이미 디스크에 있다.
        threading.Thread(target=_exit_soon, daemon=False).start()

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0") or "0")
        if not length:
            return {}
        return json.loads(self.rfile.read(length).decode("utf-8"))

    def send_text(self, body: str, content_type: str) -> None:
        data = body.encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_json(self, data: dict, status: int = 200) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def parse_iso_date(value) -> dt.date | None:
    try:
        return dt.date.fromisoformat(str(value))
    except Exception:
        return None


def resolve_run_dir(run_id: str) -> Path | None:
    if not re.fullmatch(r"\d{6}_\d{4}(?:_\d+)?", run_id or ""):
        return None
    path = (paths.RUNS_DIR / run_id).resolve()
    try:
        path.relative_to(paths.RUNS_DIR.resolve())
    except ValueError:
        return None
    return path if path.exists() and path.is_dir() else None


def same_path(left: str | Path, right: str | Path) -> bool:
    try:
        return Path(left).resolve() == Path(right).resolve()
    except Exception:
        return str(left) == str(right)


def _exit_soon(delay: float = 1.2) -> None:
    """응답이 나갈 틈을 준 뒤 프로세스를 끝낸다(업데이트 적용 전용)."""
    time.sleep(delay)
    os._exit(0)


def update_busy_reason() -> str | None:
    """지금 앱을 끊으면 잃는 게 있는가 — 있으면 그 사유.

    업데이트는 프로세스를 죽였다 살리는 일이라, 수집·보고서 생성 도중에 걸면 사용자는 몇 분치
    작업과 OpenAI 호출 비용을 잃는다. 그래서 진행 중이면 거부한다.
    """
    with JOBS_LOCK:
        running = [dict(job) for job in JOBS.values() if not job.get("done")]
    for job in running:
        if job.get("cancelled"):
            continue
        # 보고서 생성 작업만 docx 키를 갖는다(handle_generate_report_start).
        if "docx" in job:
            return "보고서 생성이 진행 중입니다. 끝난 뒤 업데이트해 주세요."
        return "자료 수집이 진행 중입니다. 끝난 뒤 업데이트해 주세요."
    return None


def main() -> int:
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    url = f"http://{HOST}:{PORT}"
    print(f"{APP_NAME} {updater.APP_VERSION}: {url}")
    # 업데이트 재시작은 브라우저를 열지 않는다 — 업데이트를 시작한 탭이 스스로 새로고침하므로
    # 여기서 또 열면 탭이 두 개가 된다(updater/apply.ps1 의 Start-App 이 이 값을 준다).
    if os.environ.get("PPCRH_NO_BROWSER") != "1":
        threading.Timer(0.4, lambda: webbrowser.open(url)).start()
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
